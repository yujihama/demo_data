# -*- coding: utf-8 -*-
"""7.source_documents 配下のインプット文書を生成するための共通ユーティリティ。

- PDF  : reportlab（IPAGothic）… 規程・方針・要領・議事録など確定版文書
- xlsx : openpyxl            … 様式・ワークシート・台帳・一覧表
- docx : python-docx         … マニュアル・手順書（改訂前提の文書）
- md   : プレーンテキスト      … 定義書・用語集
"""
from __future__ import annotations

import os
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
)

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.shared import Cm as DocxCm

# --------------------------------------------------------------------------- 共通
BASE = Path("/home/user/demo_data/7.source_documents")

COMPANY = "デモA株式会社"

_E_DIRS = {
    range(1, 10): "10_elc/01_control_environment",
    range(10, 16): "10_elc/02_risk_assessment",
    range(16, 19): "10_elc/03_control_activities",
    range(19, 25): "10_elc/04_information_communication",
    range(25, 31): "10_elc/05_monitoring",
    range(31, 35): "10_elc/06_it",
}
_F_DIRS = {
    range(1, 5): "30_close_process/01_monthly_close",
    range(5, 8): "30_close_process/02_reconciliation",
    range(8, 12): "30_close_process/03_journal_entry",
    range(12, 23): "30_close_process/04_estimates",
    range(23, 31): "30_close_process/05_consolidation",
    range(31, 33): "30_close_process/06_it_euc",
}


def doc_dir(doc_id: str) -> Path:
    """文書IDから格納先ディレクトリを解決する。"""
    area, num = doc_id.split("-")
    n = int(num)
    if area == "G":
        rel = "00_common"
    elif area == "E":
        rel = next(v for k, v in _E_DIRS.items() if n in k)
    elif area == "C":
        rel = "20_close_entity"
    elif area == "F":
        rel = next(v for k, v in _F_DIRS.items() if n in k)
    elif area == "X":
        rel = "90_records"
    else:
        raise ValueError(doc_id)
    d = BASE / rel
    d.mkdir(parents=True, exist_ok=True)
    return d


def out_path(doc_id: str, stem: str, version: str, date: str, ext: str) -> Path:
    """{ID}_{文書名}_{版数}_{施行日}.{ext}"""
    parts = [doc_id, stem]
    if version:
        parts.append(version)
    if date:
        parts.append(date)
    return doc_dir(doc_id) / ("_".join(parts) + "." + ext)


# --------------------------------------------------------------------------- PDF
_IPAG = "/usr/share/fonts/opentype/ipafont-gothic/ipag.ttf"
_IPAGP = "/usr/share/fonts/opentype/ipafont-gothic/ipagp.ttf"
pdfmetrics.registerFont(TTFont("IPAG", _IPAG))
pdfmetrics.registerFont(TTFont("IPAGP", _IPAGP))

NAVY = colors.HexColor("#1F3864")
BLUE = colors.HexColor("#2F5597")
LIGHT = colors.HexColor("#D6E4F0")
GRAY = colors.HexColor("#F2F2F2")
GREY_TXT = colors.HexColor("#595959")

_ss = getSampleStyleSheet()
PS_TITLE = ParagraphStyle("t", parent=_ss["Title"], fontName="IPAG", fontSize=17,
                          leading=23, textColor=NAVY, spaceAfter=2)
PS_SUB = ParagraphStyle("st", parent=_ss["Normal"], fontName="IPAG", fontSize=9.5,
                        leading=14, textColor=GREY_TXT, alignment=TA_CENTER, spaceAfter=10)
PS_H1 = ParagraphStyle("h1", parent=_ss["Normal"], fontName="IPAG", fontSize=12,
                       leading=17, textColor=colors.white, backColor=BLUE,
                       borderPadding=(4, 5, 4, 5), spaceBefore=10, spaceAfter=6)
PS_H2 = ParagraphStyle("h2", parent=_ss["Normal"], fontName="IPAG", fontSize=10.5,
                       leading=15, textColor=NAVY, spaceBefore=8, spaceAfter=3)
PS_BODY = ParagraphStyle("b", parent=_ss["Normal"], fontName="IPAG", fontSize=9.5,
                         leading=15, alignment=TA_LEFT, spaceAfter=3)
PS_ITEM = ParagraphStyle("i", parent=PS_BODY, leftIndent=6 * mm)
PS_NOTE = ParagraphStyle("n", parent=PS_BODY, fontSize=8.5, leading=13, textColor=GREY_TXT)
PS_CELL = ParagraphStyle("c", parent=_ss["Normal"], fontName="IPAG", fontSize=8.5, leading=12)
PS_CELLH = ParagraphStyle("ch", parent=PS_CELL, textColor=colors.white)


def _tbl(headers, rows, widths, align_center=()):
    """Paragraph 折り返し付きの表を作る。widths は mm 単位。"""
    data = [[Paragraph(str(h), PS_CELLH) for h in headers]]
    for r in rows:
        data.append([Paragraph(str(c).replace("\n", "<br/>"), PS_CELL) for c in r])
    t = Table(data, colWidths=[w * mm for w in widths], repeatRows=1)
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), BLUE),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BFBFBF")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F9FC")]),
    ]
    for c in align_center:
        style.append(("ALIGN", (c, 0), (c, -1), "CENTER"))
    t.setStyle(TableStyle(style))
    return t


def _meta_table(pairs, widths=(32, 58, 32, 58)):
    data = []
    it = list(pairs)
    for i in range(0, len(it), 2):
        chunk = it[i:i + 2]
        row = []
        for k, v in chunk:
            row += [Paragraph(str(k), PS_CELL), Paragraph(str(v), PS_CELL)]
        while len(row) < 4:
            row.append(Paragraph("", PS_CELL))
        data.append(row)
    t = Table(data, colWidths=[w * mm for w in widths])
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BFBFBF")),
        ("BACKGROUND", (0, 0), (0, -1), LIGHT),
        ("BACKGROUND", (2, 0), (2, -1), LIGHT),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return t


def _footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("IPAG", 7.5)
    canvas.setFillColor(GREY_TXT)
    w = doc.pagesize[0]
    canvas.drawString(20 * mm, 10 * mm, getattr(doc, "_footer_left", ""))
    canvas.drawRightString(w - 20 * mm, 10 * mm, f"- {doc.page} -")
    canvas.restoreState()


def build_pdf(path: Path, title: str, subtitle: str, meta_pairs, blocks,
              footer_left: str = "", landscape_mode: bool = False):
    """blocks: ('h1'|'h2'|'p'|'item'|'note', text) / ('table', headers, rows, widths) / ('spacer', mm) / ('pagebreak',)"""
    size = landscape(A4) if landscape_mode else A4
    doc = SimpleDocTemplate(
        str(path), pagesize=size,
        leftMargin=20 * mm, rightMargin=20 * mm, topMargin=18 * mm, bottomMargin=18 * mm,
        title=title, author=COMPANY,
    )
    doc._footer_left = footer_left or f"{COMPANY}　{title}"
    story = [Paragraph(title, PS_TITLE)]
    if subtitle:
        story.append(Paragraph(subtitle, PS_SUB))
    else:
        story.append(Spacer(1, 6))
    if meta_pairs:
        story.append(_meta_table(meta_pairs))
        story.append(Spacer(1, 8))
    for b in blocks:
        kind = b[0]
        if kind == "h1":
            story.append(Paragraph(b[1], PS_H1))
        elif kind == "h2":
            story.append(Paragraph(b[1], PS_H2))
        elif kind == "p":
            story.append(Paragraph(b[1], PS_BODY))
        elif kind == "item":
            story.append(Paragraph(b[1], PS_ITEM))
        elif kind == "note":
            story.append(Paragraph(b[1], PS_NOTE))
        elif kind == "table":
            _, headers, rows, widths = b[:4]
            ac = b[4] if len(b) > 4 else ()
            story.append(Spacer(1, 2))
            story.append(_tbl(headers, rows, widths, ac))
            story.append(Spacer(1, 5))
        elif kind == "spacer":
            story.append(Spacer(1, b[1]))
        elif kind == "pagebreak":
            story.append(PageBreak())
    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return path


def build_regulation(path: Path, doc_no: str, title: str, meta_pairs, purpose_note,
                     chapters, history, appendix=None, footer_left=""):
    """規程・要領の標準体裁（章→条→項）でPDFを生成する。

    chapters: [(章名, [(条見出し, [項... / ('table',headers,rows,widths)]), ...]), ...]
    history : [(版数, 改訂日, 改訂内容, 承認者), ...]
    appendix: [(別表名, headers, rows, widths), ...]
    """
    blocks = []
    if purpose_note:
        blocks.append(("note", purpose_note))
        blocks.append(("spacer", 4))
    for ch_name, articles in chapters:
        blocks.append(("h1", ch_name))
        for art_title, items in articles:
            blocks.append(("h2", art_title))
            n = 0
            for it in items:
                if isinstance(it, tuple) and it and it[0] == "table":
                    blocks.append(("table",) + tuple(it[1:]))
                elif isinstance(it, tuple) and it and it[0] == "raw":
                    blocks.append(("item", it[1]))
                else:
                    n += 1
                    prefix = "" if len(items) == 1 else f"{n}　"
                    blocks.append(("p", prefix + str(it)))
    if appendix:
        blocks.append(("pagebreak",))
        for name, headers, rows, widths in appendix:
            blocks.append(("h1", name))
            blocks.append(("table", headers, rows, widths))
    if history:
        blocks.append(("h1", "改訂履歴"))
        blocks.append(("table", ["版数", "改訂日", "改訂内容", "承認"],
                       history, [18, 25, 95, 32], (0, 1, 3)))
    return build_pdf(path, title, f"文書番号：{doc_no}　／　{COMPANY}",
                     meta_pairs, blocks, footer_left or f"{COMPANY}　{doc_no} {title}")


# --------------------------------------------------------------------------- xlsx
X_NAVY, X_BLUE, X_LIGHT, X_GRAY, X_YELLOW, X_ORANGE = (
    "1F3864", "2F5597", "D6E4F0", "F2F2F2", "FFF2CC", "FCE4D6")
_THIN = Side(style="thin", color="BFBFBF")
XB = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
XF_TITLE = Font(name="Yu Gothic", size=14, bold=True, color=X_NAVY)
XF_SUB = Font(name="Yu Gothic", size=9, color="595959")
XF_HEAD = Font(name="Yu Gothic", size=9, bold=True, color="FFFFFF")
XF_BODY = Font(name="Yu Gothic", size=9)
XF_BOLD = Font(name="Yu Gothic", size=9, bold=True)
XF_SMALL = Font(name="Yu Gothic", size=8, color="595959")
XFILL_HEAD = PatternFill("solid", fgColor=X_BLUE)
XFILL_SEC = PatternFill("solid", fgColor=X_LIGHT)
XFILL_TOT = PatternFill("solid", fgColor=X_GRAY)
XFILL_IN = PatternFill("solid", fgColor=X_YELLOW)
XFILL_WARN = PatternFill("solid", fgColor=X_ORANGE)
XNUM = "#,##0;[Red]△#,##0"


def x_widths(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w


def x_title(ws, row, text, subtitle=None):
    c = ws.cell(row=row, column=1, value=text)
    c.font = XF_TITLE
    row += 1
    if subtitle:
        s = ws.cell(row=row, column=1, value=subtitle)
        s.font = XF_SUB
        row += 1
    return row + 1


def x_meta(ws, row, pairs, val_span=4):
    for k, v in pairs:
        ck = ws.cell(row=row, column=1, value=k)
        ck.font = XF_BOLD
        ck.fill = XFILL_SEC
        ck.border = XB
        cv = ws.cell(row=row, column=2, value=v)
        cv.font = XF_BODY
        cv.border = XB
        cv.alignment = Alignment(wrap_text=True, vertical="top")
        row += 1
    return row + 1


def x_head(ws, row, headers, fill=XFILL_HEAD):
    for i, h in enumerate(headers, 1):
        c = ws.cell(row=row, column=i, value=h)
        c.font = XF_HEAD
        c.fill = fill
        c.border = XB
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[row].height = 30
    return row + 1


def x_row(ws, row, values, numcols=(), bold=False, fill=None, wrap=True):
    for i, v in enumerate(values, 1):
        c = ws.cell(row=row, column=i, value=v)
        c.font = XF_BOLD if bold else XF_BODY
        c.border = XB
        c.alignment = Alignment(wrap_text=wrap, vertical="top")
        if fill:
            c.fill = fill
        if i in numcols:
            c.number_format = XNUM
            c.alignment = Alignment(horizontal="right", vertical="top")
    return row + 1


def x_note(ws, row, text):
    c = ws.cell(row=row, column=1, value=text)
    c.font = XF_SMALL
    c.alignment = Alignment(wrap_text=True, vertical="top")
    return row + 1


def x_section(ws, row, text, ncols=8):
    for i in range(1, ncols + 1):
        c = ws.cell(row=row, column=i, value=text if i == 1 else None)
        c.fill = XFILL_SEC
        c.border = XB
        if i == 1:
            c.font = Font(name="Yu Gothic", size=10, bold=True, color=X_NAVY)
    return row + 1


def build_xlsx(path: Path, sheets):
    """sheets: [(sheet_name, build_fn(ws)), ...]"""
    wb = Workbook()
    wb.remove(wb.active)
    for name, fn in sheets:
        ws = wb.create_sheet(title=name[:31])
        fn(ws)
    wb.save(str(path))
    return path


def simple_form(path: Path, sheet_name: str, title: str, subtitle: str,
                meta_pairs, headers, widths, rows, notes=(), blank_rows=0,
                numcols=(), freeze="A1"):
    """単票の様式・台帳を1シートで作る汎用ビルダ。"""
    def build(ws):
        x_widths(ws, widths)
        r = x_title(ws, 1, title, subtitle)
        if meta_pairs:
            r = x_meta(ws, r, meta_pairs)
        r = x_head(ws, r, headers)
        head_row = r - 1
        for row_vals in rows:
            r = x_row(ws, r, row_vals, numcols=numcols)
        for _ in range(blank_rows):
            r = x_row(ws, r, [None] * len(headers), fill=XFILL_IN)
        r += 1
        for n in notes:
            r = x_note(ws, r, n)
        ws.freeze_panes = ws.cell(row=head_row + 1, column=1)
        ws.auto_filter.ref = (f"A{head_row}:{get_column_letter(len(headers))}"
                              f"{head_row + len(rows) + blank_rows}")
    return build_xlsx(path, [(sheet_name, build)])


# --------------------------------------------------------------------------- docx
def _set_jp_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Yu Gothic"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")


class DocxBuilder:
    def __init__(self, title, subtitle=None):
        self.d = docx.Document()
        st = self.d.styles["Normal"]
        st.font.name = "Yu Gothic"
        st.font.size = Pt(10)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        for s in self.d.sections:
            s.top_margin = DocxCm(2.0)
            s.bottom_margin = DocxCm(2.0)
            s.left_margin = DocxCm(2.0)
            s.right_margin = DocxCm(2.0)
        p = self.d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_jp_font(p.add_run(title), 16, True, "1F3864")
        if subtitle:
            q = self.d.add_paragraph()
            q.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_jp_font(q.add_run(subtitle), 9, False, "595959")

    def meta(self, pairs):
        t = self.d.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for k, v in pairs:
            cells = t.add_row().cells
            _set_jp_font(cells[0].paragraphs[0].add_run(str(k)), 9, True)
            _set_jp_font(cells[1].paragraphs[0].add_run(str(v)), 9)
        self.d.add_paragraph()
        return self

    def h1(self, text):
        p = self.d.add_paragraph()
        _set_jp_font(p.add_run(text), 13, True, "2F5597")
        return self

    def h2(self, text):
        p = self.d.add_paragraph()
        _set_jp_font(p.add_run(text), 11, True, "1F3864")
        return self

    def p(self, text, size=10, indent=None):
        p = self.d.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = DocxCm(indent)
        _set_jp_font(p.add_run(text), size)
        return self

    def bullets(self, items):
        for it in items:
            p = self.d.add_paragraph(style="List Bullet")
            _set_jp_font(p.add_run(str(it)), 10)
        return self

    def note(self, text):
        p = self.d.add_paragraph()
        _set_jp_font(p.add_run(text), 8.5, False, "595959")
        return self

    def table(self, headers, rows, widths=None):
        t = self.d.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_jp_font(cell.paragraphs[0].add_run(str(h)), 8.5, True)
        for r in rows:
            cells = t.add_row().cells
            for i, v in enumerate(r):
                _set_jp_font(cells[i].paragraphs[0].add_run("" if v is None else str(v)), 8.5)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = DocxCm(w)
        self.d.add_paragraph()
        return self

    def pagebreak(self):
        self.d.add_page_break()
        return self

    def save(self, path):
        self.d.save(str(path))
        return path


# --------------------------------------------------------------------------- md
def write_md(path: Path, text: str):
    path.write_text(text.rstrip() + "\n", encoding="utf-8")
    return path


# --------------------------------------------------------------------------- 進捗
_GENERATED = []


def done(path):
    rel = Path(path).relative_to(BASE)
    _GENERATED.append(str(rel))
    return path


def report():
    print(f"generated: {len(_GENERATED)} files")
    for p in _GENERATED:
        print("  ", p)
